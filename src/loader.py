"""
loader.py - Text extraction for local documents.
"""

import os

import docx
import pypdf


def extract_text_from_pdf(file_path_or_obj) -> str:
    """Extract text from a PDF file path or file-like object."""
    try:
        reader = pypdf.PdfReader(file_path_or_obj)
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n".join(text)
    except Exception as exc:
        raise RuntimeError(f"Failed to extract text from PDF: {exc}") from exc


def extract_text_from_docx(file_path_or_obj) -> str:
    """Extract paragraphs and table text from a Word document."""
    try:
        document = docx.Document(file_path_or_obj)
        text = []

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())

        for table in document.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))

        return "\n".join(text)
    except Exception as exc:
        raise RuntimeError(f"Failed to extract text from Word document: {exc}") from exc


def load_single_document(file_path_or_obj, filename: str) -> str:
    """Load text from a .txt, .md, .pdf, or .docx file."""
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".txt", ".md"):
        if isinstance(file_path_or_obj, str):
            with open(file_path_or_obj, "r", encoding="utf-8", errors="ignore") as file:
                return file.read().strip()

        content = file_path_or_obj.read()
        if isinstance(content, bytes):
            return content.decode("utf-8", errors="ignore").strip()
        return str(content).strip()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path_or_obj)

    if ext == ".docx":
        return extract_text_from_docx(file_path_or_obj)

    raise ValueError(f"Unsupported file format: {ext}")
