"""
Basic tests for the local RAG pipeline.
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest

from src.retriever import cosine_similarity, format_context


def test_cosine_similarity_identical():
    v = [1.0, 0.5, -0.3]
    assert abs(cosine_similarity(v, v) - 1.0) < 1e-6


def test_cosine_similarity_orthogonal():
    a = [1.0, 0.0]
    b = [0.0, 1.0]
    assert abs(cosine_similarity(a, b)) < 1e-6


def test_cosine_similarity_zero_vector():
    assert cosine_similarity([0.0, 0.0], [1.0, 0.5]) == 0.0


def test_database_insert_and_retrieve(tmp_path, monkeypatch):
    import src.config as cfg

    monkeypatch.setattr(cfg, "DATABASE_PATH", str(tmp_path / "test.db"))

    import importlib
    import src.database as db_module

    importlib.reload(db_module)
    db_module.create_tables()
    db_module.clear_database()

    db_module.insert_chunk("test.txt", "Hello world", [0.1, 0.2, 0.3])
    assert db_module.count_chunks() == 1

    chunks = db_module.get_all_chunks()
    assert chunks[0]["content"] == "Hello world"
    assert chunks[0]["source"] == "test.txt"
    assert len(chunks[0]["embedding"]) == 3


def test_format_context_empty():
    result = format_context([])
    assert "No relevant" in result


def test_format_context_with_chunks():
    chunks = [
        {"source": "faq.txt", "content": "Python is great.", "score": 0.9},
        {"source": "manual.txt", "content": "Install with pip.", "score": 0.7},
    ]
    result = format_context(chunks)
    assert "faq.txt" in result
    assert "Python is great." in result
    assert "manual.txt" in result


def test_pipeline_returns_dict(tmp_path, monkeypatch):
    from src.foundry_runtime import get_runtime_status

    if not get_runtime_status()["foundry_ready"]:
        pytest.skip("Foundry Local models are not cached")

    import src.config as cfg

    monkeypatch.setattr(cfg, "DATABASE_PATH", str(tmp_path / "integration_test.db"))

    import importlib
    import src.database as db_module

    importlib.reload(db_module)
    db_module.create_tables()
    db_module.clear_database()

    from src.embedder import embed_text

    text = "RAG stands for Retrieval-Augmented Generation."
    emb = embed_text(text)
    db_module.insert_chunk("test.txt", text, emb)

    import src.pipeline as pipeline_module

    importlib.reload(pipeline_module)
    result = pipeline_module.ask("What does RAG stand for?")

    assert isinstance(result, dict)
    assert "answer" in result
    assert "sources" in result
    assert "chunks" in result
    assert result["sources"] == ["test.txt"]


def test_docx_loading(tmp_path):
    import docx
    from src.loader import load_single_document

    doc_path = tmp_path / "test.docx"
    doc = docx.Document()
    doc.add_paragraph("This is a paragraph in a Word document.")
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Passport123"
    doc.save(str(doc_path))

    text = load_single_document(str(doc_path), "test.docx")
    assert "This is a paragraph in a Word document." in text
    assert "ID | Passport123" in text


def test_md_loading(tmp_path):
    from src.loader import load_single_document

    md_path = tmp_path / "test.md"
    md_path.write_text("# Heading\nThis is a markdown file content.", encoding="utf-8")

    text = load_single_document(str(md_path), "test.md")
    assert "# Heading" in text
    assert "This is a markdown file content." in text


def test_sentence_aware_chunking():
    from scripts.ingest import split_into_chunks

    text = "First sentence about Python. Second sentence is longer. Third sentence details SQLite."
    # With a chunk_size that fits two sentences but not three
    chunks = split_into_chunks(text, chunk_size=65, overlap=10)
    
    assert len(chunks) >= 2
    assert "First sentence about Python." in chunks[0]
    assert "Second sentence is longer." in chunks[0]
    assert "Third sentence details SQLite." in chunks[1]
    # Check that sentences are not cut in half
    for c in chunks:
        assert c.strip().endswith(".")


def test_chunking_enforces_hard_maximum_for_long_text():
    from scripts.ingest import split_into_chunks

    text = "word " * 500
    chunks = split_into_chunks(text, chunk_size=120, overlap=20)

    assert len(chunks) > 1
    assert all(0 < len(chunk) <= 120 for chunk in chunks)


def test_chunking_rejects_invalid_settings():
    from scripts.ingest import split_into_chunks

    with pytest.raises(ValueError):
        split_into_chunks("text", chunk_size=100, overlap=100)


def test_retrieval_confidence_gate():
    from src.retriever import _passes_confidence_gate

    assert _passes_confidence_gate(
        {"score": 0.55, "semantic_score": 0.60, "keyword_score": 0.0}
    )
    assert _passes_confidence_gate(
        {"score": 0.45, "semantic_score": 0.42, "keyword_score": 0.50}
    )
    assert not _passes_confidence_gate(
        {"score": 0.28, "semantic_score": 0.40, "keyword_score": 0.0}
    )


def test_pipeline_refuses_when_retrieval_has_no_evidence(monkeypatch):
    import src.pipeline as pipeline

    monkeypatch.setattr(pipeline, "count_chunks", lambda: 1)
    monkeypatch.setattr(pipeline, "find_relevant_chunks", lambda *args, **kwargs: [])

    def should_not_run(*args, **kwargs):
        raise AssertionError("LLM should not run without relevant context")

    monkeypatch.setattr(pipeline, "generate_answer", should_not_run)
    result = pipeline.ask("What is the capital of Brazil?")

    assert result["answer"] == pipeline.NO_INFORMATION_ANSWER
    assert result["sources"] == []
    assert result["chunks"] == []


def test_empty_document_scope_does_not_search_all_documents(monkeypatch):
    import src.pipeline as pipeline

    monkeypatch.setattr(pipeline, "count_chunks", lambda: 1)
    observed = {}

    def fake_retrieval(question, allowed_sources=None):
        observed["allowed_sources"] = allowed_sources
        return []

    monkeypatch.setattr(pipeline, "find_relevant_chunks", fake_retrieval)
    result = pipeline.ask("Any question", preferred_sources=[])

    assert observed["allowed_sources"] == []
    assert result["answer"] == pipeline.NO_INFORMATION_ANSWER


def test_refusal_cannot_append_outside_knowledge():
    from src.pipeline import NO_INFORMATION_ANSWER, _sanitize_grounded_answer

    unsafe = (
        "I don't have that information in my documents. "
        "However, the capital of Brazil is Brasilia."
    )
    assert _sanitize_grounded_answer(unsafe) == NO_INFORMATION_ANSWER
